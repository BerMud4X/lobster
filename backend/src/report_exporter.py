from pathlib import Path
from logger import logger


def export_docx(sections: list[dict], figures: dict, output_path: Path, title: str = "Clinical Report") -> Path:
    """
    Exports a report to .docx.
    sections: list of {"heading": str, "body": str}
    figures: {name: (svg_path, png_path)}
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Sections
    for section in sections:
        if section.get("heading"):
            doc.add_heading(section["heading"], level=1)
        if section.get("body"):
            p = doc.add_paragraph(section["body"])
            p.paragraph_format.space_after = Pt(8)

    # Figures
    if figures:
        doc.add_heading("Figures", level=1)
        for name, (svg_path, png_path) in figures.items():
            if png_path and Path(png_path).exists():
                doc.add_picture(str(png_path), width=Inches(5.5))
                cap = doc.add_paragraph(f"Figure — {name.replace('_', ' ').title()}")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
                doc.add_paragraph()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info(f"[Exporter] .docx saved: {output_path}")
    return output_path


def _get_dejavu_fonts() -> tuple[str, str, str]:
    """Returns paths to DejaVuSans regular, bold, and italic TTF files."""
    import matplotlib
    import os
    font_dir = os.path.join(os.path.dirname(matplotlib.__file__), "mpl-data", "fonts", "ttf")
    return (
        os.path.join(font_dir, "DejaVuSans.ttf"),
        os.path.join(font_dir, "DejaVuSans-Bold.ttf"),
        os.path.join(font_dir, "DejaVuSans-Oblique.ttf"),
    )


def export_pdf(sections: list[dict], figures: dict, output_path: Path, title: str = "Clinical Report") -> Path:
    """
    Exports a report to .pdf using fpdf2 with Unicode (DejaVu) font support.
    sections: list of {"heading": str, "body": str}
    figures: {name: (svg_path, png_path)}
    """
    from fpdf import FPDF

    regular, bold, italic = _get_dejavu_fonts()

    pdf = FPDF()
    pdf.add_font("DejaVu", "", regular)
    pdf.add_font("DejaVu", "B", bold)
    pdf.add_font("DejaVu", "I", italic)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("DejaVu", "B", 18)
    pdf.set_text_color(45, 106, 159)
    pdf.multi_cell(0, 12, title, align="C")
    pdf.ln(6)

    # Sections
    for section in sections:
        if section.get("heading"):
            pdf.set_font("DejaVu", "B", 13)
            pdf.set_text_color(45, 106, 159)
            pdf.ln(4)
            pdf.multi_cell(0, 8, section["heading"])
            pdf.ln(2)

        if section.get("body"):
            pdf.set_font("DejaVu", "", 10)
            pdf.set_text_color(30, 30, 30)
            pdf.multi_cell(0, 6, section["body"])
            pdf.ln(3)

    # Figures
    if figures:
        pdf.ln(4)
        pdf.set_font("DejaVu", "B", 13)
        pdf.set_text_color(45, 106, 159)
        pdf.multi_cell(0, 8, "Figures")
        pdf.ln(2)

        for name, (svg_path, png_path) in figures.items():
            if png_path and Path(png_path).exists():
                pdf.image(str(png_path), w=170)
                pdf.set_font("DejaVu", "I", 9)
                pdf.set_text_color(80, 80, 80)
                pdf.multi_cell(0, 6, f"Figure — {name.replace('_', ' ').title()}", align="C")
                pdf.ln(4)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(output_path))
    logger.info(f"[Exporter] .pdf saved: {output_path}")
    return output_path


def export_report(
    sections: list[dict],
    figures: dict,
    output_path: str,
    fmt: str = "pdf",
    title: str = "Clinical Report",
) -> Path:
    """
    Main export entry point.
    fmt: 'pdf' or 'docx'
    """
    path = Path(output_path)
    if not path.suffix:
        path = path.with_suffix(f".{fmt}")

    if fmt == "docx":
        return export_docx(sections, figures, path, title=title)
    else:
        return export_pdf(sections, figures, path, title=title)
